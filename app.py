from flask import Flask, render_template, request, send_file
from docx import Document
import os
from datetime import datetime
import zipfile
import tempfile
import json

app = Flask(__name__)

from docx.oxml.ns import qn
from docx.shared import Pt

def set_aptos_font(paragraphs):
    for para in paragraphs:
        for run in para.runs:
            run.font.name = 'Aptos'
            run.font.size = Pt(12)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')

def update_word_file(promoter_name, project_name, registration_no):
    template_path = os.path.join('resources', 'complaintDeclaration.docx')
    output_folder = os.path.join('generated', 'complaintDeclaration')
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(template_path)
    replacements = {'{{promoter_name}}': promoter_name, '{{project_name}}': project_name, '{{registration_no}}': registration_no}
    def replace_text_with_bold(paragraphs, replacements):
        for para in paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            if not any(ph in full_text for ph in replacements.keys()):
                continue
            new_runs = [(full_text, False)]
            for placeholder, value in replacements.items():
                updated_runs = []
                for text, bold in new_runs:
                    if placeholder in text:
                        parts = text.split(placeholder)
                        for i, part in enumerate(parts):
                            updated_runs.append((part, bold))
                            if i < len(parts) - 1:
                                updated_runs.append((value, True))
                    else:
                        updated_runs.append((text, bold))
                new_runs = updated_runs
            para.clear()
            for text, bold in new_runs:
                run = para.add_run(text)
                run.bold = bold
    replace_text_with_bold(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_with_bold(cell.paragraphs, replacements)
    set_aptos_font(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_aptos_font(cell.paragraphs)
    sanitized_name = promoter_name.replace(" ", "_")
    filename = f"{sanitized_name}complaintDeclaration.docx"
    output_path = os.path.join(output_folder, filename)
    doc.save(output_path)
    return output_path

def generate_no_complaint_file(promoter_name, project_name, registration_no, date):
    template_path = os.path.join('resources', 'noComplaintsDeclaration.docx')
    output_folder = os.path.join('generated', 'noComplaintsDeclaration')
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(template_path)
    date_fmt = datetime.strptime(date, "%Y-%m-%d").strftime("%d-%m-%Y")
    replacements = {'{{promoter_name}}': promoter_name, '{{project_name}}': project_name, '{{registration_no}}': registration_no, '{{date}}': date_fmt}
    def replace_text_with_bold(paragraphs, replacements):
        for para in paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            if not any(ph in full_text for ph in replacements.keys()):
                continue
            new_runs = [(full_text, False)]
            for placeholder, value in replacements.items():
                updated_runs = []
                for text, bold in new_runs:
                    if placeholder in text:
                        parts = text.split(placeholder)
                        for i, part in enumerate(parts):
                            updated_runs.append((part, bold))
                            if i < len(parts) - 1:
                                updated_runs.append((value, True))
                    else:
                        updated_runs.append((text, bold))
                new_runs = updated_runs
            para.clear()
            for text, bold in new_runs:
                run = para.add_run(text)
                run.bold = bold
    replace_text_with_bold(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_with_bold(cell.paragraphs, replacements)
    set_aptos_font(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_aptos_font(cell.paragraphs)
    sanitized_name = promoter_name.replace(" ", "_")
    filename = f"{sanitized_name}_noComplaintsDeclaration.docx"
    output_path = os.path.join(output_folder, filename)
    doc.save(output_path)
    return output_path

def update_extension_application(date, registration_no, extension_date, promoter_name):
    template_path = os.path.join('resources', 'Extension Application under Section 7(3).docx')
    output_folder = os.path.join('generated', 'extension')
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(template_path)
    date_fmt = datetime.strptime(date, "%Y-%m-%d").strftime("%d-%m-%Y")
    extension_date_fmt = datetime.strptime(extension_date, "%Y-%m-%d").strftime("%d-%m-%Y")
    replacements = {'{{promoter_name}}': promoter_name, '{{extension_date}}': extension_date_fmt, '{{registration_no}}': registration_no, '{{date}}': date_fmt}
    def replace_text_with_bold(paragraphs, replacements):
        for para in paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            if not any(ph in full_text for ph in replacements.keys()):
                continue
            new_runs = [(full_text, False)]
            for placeholder, value in replacements.items():
                updated_runs = []
                for text, bold in new_runs:
                    if placeholder in text:
                        parts = text.split(placeholder)
                        for i, part in enumerate(parts):
                            updated_runs.append((part, bold))
                            if i < len(parts) - 1:
                                updated_runs.append((value, True))
                    else:
                        updated_runs.append((text, bold))
                new_runs = updated_runs
            para.clear()
            for text, bold in new_runs:
                run = para.add_run(text)
                run.bold = bold
    replace_text_with_bold(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_with_bold(cell.paragraphs, replacements)
    set_aptos_font(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_aptos_font(cell.paragraphs)
    sanitized_name = promoter_name.replace(" ", "_")
    filename = f"{sanitized_name}_Extension Application.docx"
    output_path = os.path.join(output_folder, filename)
    doc.save(output_path)
    return output_path

def update_project_pert_file(project_name, registration_no, extension_date, promoter_name):
    template_path = os.path.join('resources', 'projectPert.docx')
    output_folder = os.path.join('generated', 'project_pert')
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(template_path)
    extension_date_fmt = datetime.strptime(extension_date, "%Y-%m-%d").strftime("%d-%m-%Y")
    replacements = {'{{promoter_name}}': promoter_name, '{{extension_date}}': extension_date_fmt, '{{registration_no}}': registration_no, '{{project_name}}': project_name}
    def replace_text_with_bold(paragraphs, replacements):
        for para in paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            if not any(ph in full_text for ph in replacements.keys()):
                continue
            new_runs = [(full_text, False)]
            for placeholder, value in replacements.items():
                updated_runs = []
                for text, bold in new_runs:
                    if placeholder in text:
                        parts = text.split(placeholder)
                        for i, part in enumerate(parts):
                            updated_runs.append((part, bold))
                            if i < len(parts) - 1:
                                updated_runs.append((value, True))
                    else:
                        updated_runs.append((text, bold))
                new_runs = updated_runs
            para.clear()
            for text, bold in new_runs:
                run = para.add_run(text)
                run.bold = bold
    replace_text_with_bold(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_with_bold(cell.paragraphs, replacements)
    set_aptos_font(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_aptos_font(cell.paragraphs)
    sanitized_name = promoter_name.replace(" ", "_")
    filename = f"{sanitized_name}_projectPert.docx"
    output_path = os.path.join(output_folder, filename)
    doc.save(output_path)
    return output_path

def update_cersai_file(promoter_name, project_name, office_address, project_address, date):
    template_path = os.path.join('resources','CERSAI Declaration.docx')
    output_folder = os.path.join('generated','CERSAI_Declaration')
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(template_path)
    date_fmt = datetime.strptime(date, "%Y-%m-%d").strftime("%d-%m-%Y")
    replacements = {'{{promoter_name}}': promoter_name, '{{project_name}}': project_name, '{{office_address}}': office_address, '{{project_address}}': project_address, '{{date}}': date_fmt}
    def replace_text_with_bold(paragraphs, replacements):
        for para in paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            if not any(ph in full_text for ph in replacements.keys()):
                continue
            new_runs = [(full_text, False)]
            for placeholder, value in replacements.items():
                updated_runs = []
                for text, bold in new_runs:
                    if placeholder in text:
                        parts = text.split(placeholder)
                        for i, part in enumerate(parts):
                            updated_runs.append((part, bold))
                            if i < len(parts) - 1:
                                updated_runs.append((value, True))
                    else:
                        updated_runs.append((text, bold))
                new_runs = updated_runs
            para.clear()
            for text, bold in new_runs:
                run = para.add_run(text)
                run.bold = bold
    replace_text_with_bold(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_with_bold(cell.paragraphs, replacements)
    set_aptos_font(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_aptos_font(cell.paragraphs)
    sanitized_name = promoter_name.replace(" ", "_")
    filename = f"{sanitized_name}_CERSAI DECLARATION.docx"
    output_path = os.path.join(output_folder, filename)
    doc.save(output_path)
    return output_path

def update_authorization_letter(promoter_name, project_name, date):
    template_path = os.path.join('resources', 'AUTHORIZATION LETTER.docx')
    output_folder = os.path.join('generated', 'Authorization')
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(template_path)
    date_fmt = datetime.strptime(date, "%Y-%m-%d").strftime("%d-%m-%Y")
    replacements = {'{{promoter_name}}': promoter_name, '{{project_name}}': project_name, '{{date}}': date_fmt}
    def replace_text_with_bold(paragraphs, replacements):
        for para in paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            if not any(ph in full_text for ph in replacements.keys()):
                continue
            new_runs = [(full_text, False)]
            for placeholder, value in replacements.items():
                updated_runs = []
                for text, bold in new_runs:
                    if placeholder in text:
                        parts = text.split(placeholder)
                        for i, part in enumerate(parts):
                            updated_runs.append((part, bold))
                            if i < len(parts) - 1:
                                updated_runs.append((value, True))
                    else:
                        updated_runs.append((text, bold))
                new_runs = updated_runs
            para.clear()
            for text, bold in new_runs:
                run = para.add_run(text)
                run.bold = bold
    replace_text_with_bold(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_with_bold(cell.paragraphs, replacements)
    set_aptos_font(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_aptos_font(cell.paragraphs)
    sanitized_name = promoter_name.replace(" ", "_")
    filename = f"{sanitized_name}_Authorization_Letter.docx"
    output_path = os.path.join(output_folder, filename)
    doc.save(output_path)
    return output_path

def update_annexure_a(promoter_name, date):
    template_path = os.path.join('resources', 'Annexure A.docx')
    output_folder = os.path.join('generated', 'Annexure_A')
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(template_path)
    date_fmt = datetime.strptime(date, "%Y-%m-%d").strftime("%d-%m-%Y")
    replacements = {'{{promoter_name}}': promoter_name, '{{date}}': date_fmt}
    def replace_text_with_bold(paragraphs, replacements):
        for para in paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            if not any(ph in full_text for ph in replacements.keys()):
                continue
            new_runs = [(full_text, False)]
            for placeholder, value in replacements.items():
                updated_runs = []
                for text, bold in new_runs:
                    if placeholder in text:
                        parts = text.split(placeholder)
                        for i, part in enumerate(parts):
                            updated_runs.append((part, bold))
                            if i < len(parts) - 1:
                                updated_runs.append((value, True))
                    else:
                        updated_runs.append((text, bold))
                new_runs = updated_runs
            para.clear()
            for text, bold in new_runs:
                run = para.add_run(text)
                run.bold = bold
    replace_text_with_bold(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_with_bold(cell.paragraphs, replacements)
    set_aptos_font(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_aptos_font(cell.paragraphs)
    sanitized_name = promoter_name.replace(" ", "_")
    filename = f"{sanitized_name}_Annexure_a.docx"
    output_path = os.path.join(output_folder, filename)
    doc.save(output_path)
    return output_path

def update_AffidavitReason_for_Extension(promoter_name, project_name, registration_no):
    template_path = os.path.join('resources', 'Affidavit Reason for Extension.docx')
    output_folder = os.path.join('generated', 'Affidavit')
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(template_path)
    replacements = {'{{promoter_name}}': promoter_name, '{{project_name}}': project_name, '{{registration_no}}': registration_no}
    def replace_text_with_bold(paragraphs, replacements):
        for para in paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            if not any(ph in full_text for ph in replacements.keys()):
                continue
            new_runs = [(full_text, False)]
            for placeholder, value in replacements.items():
                updated_runs = []
                for text, bold in new_runs:
                    if placeholder in text:
                        parts = text.split(placeholder)
                        for i, part in enumerate(parts):
                            updated_runs.append((part, bold))
                            if i < len(parts) - 1:
                                updated_runs.append((value, True))
                    else:
                        updated_runs.append((text, bold))
                new_runs = updated_runs
            para.clear()
            for text, bold in new_runs:
                run = para.add_run(text)
                run.bold = bold
    replace_text_with_bold(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_with_bold(cell.paragraphs, replacements)
    set_aptos_font(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_aptos_font(cell.paragraphs)
    sanitized_name = promoter_name.replace(" ", "_")
    filename = f"{sanitized_name}_AffidavitReason_for_Extension.docx"
    output_path = os.path.join(output_folder, filename)
    doc.save(output_path)
    return output_path

def update_consent_Extension_Tabular(promoter_name, project_name, registration_no, extension_date):
    template_path = os.path.join('resources', 'Consent for Extension-Tabular.docx')
    output_folder = os.path.join('generated', 'Consent for Extension')
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(template_path)
    extension_date_fmt = datetime.strptime(extension_date, "%Y-%m-%d").strftime("%d-%m-%Y")
    replacements = {'{{promoter_name}}': promoter_name, '{{project_name}}': project_name, '{{registration_no}}': registration_no, '{{extension_date}}': extension_date_fmt}
    def replace_text_with_bold(paragraphs, replacements):
        for para in paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            if not any(ph in full_text for ph in replacements.keys()):
                continue
            new_runs = [(full_text, False)]
            for placeholder, value in replacements.items():
                updated_runs = []
                for text, bold in new_runs:
                    if placeholder in text:
                        parts = text.split(placeholder)
                        for i, part in enumerate(parts):
                            updated_runs.append((part, bold))
                            if i < len(parts) - 1:
                                updated_runs.append((value, True))
                    else:
                        updated_runs.append((text, bold))
                new_runs = updated_runs
            para.clear()
            for text, bold in new_runs:
                run = para.add_run(text)
                run.bold = bold
    replace_text_with_bold(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_with_bold(cell.paragraphs, replacements)
    set_aptos_font(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_aptos_font(cell.paragraphs)
    sanitized_name = promoter_name.replace(" ", "_")
    filename = f"{sanitized_name}_Consent for Extension.docx"
    output_path = os.path.join(output_folder, filename)
    doc.save(output_path)
    return output_path

def update_Declaration_For_Extension(promoter_name, project_name, registration_no, extension_date):
    template_path = os.path.join('resources', 'Declaration For Extension.docx')
    output_folder = os.path.join('generated', 'DeclarationExtension')
    os.makedirs(output_folder, exist_ok=True)
    extension_date_fmt = datetime.strptime(extension_date, "%Y-%m-%d").strftime("%d-%m-%Y")
    doc = Document(template_path)
    replacements = {'{{promoter_name}}': promoter_name, '{{project_name}}': project_name, '{{registration_no}}': registration_no, '{{extension_date}}': extension_date_fmt}
    def replace_text_with_bold(paragraphs, replacements):
        for para in paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            if not any(ph in full_text for ph in replacements.keys()):
                continue
            new_runs = [(full_text, False)]
            for placeholder, value in replacements.items():
                updated_runs = []
                for text, bold in new_runs:
                    if placeholder in text:
                        parts = text.split(placeholder)
                        for i, part in enumerate(parts):
                            updated_runs.append((part, bold))
                            if i < len(parts) - 1:
                                updated_runs.append((value, True))
                    else:
                        updated_runs.append((text, bold))
                new_runs = updated_runs
            para.clear()
            for text, bold in new_runs:
                run = para.add_run(text)
                run.bold = bold
    replace_text_with_bold(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_with_bold(cell.paragraphs, replacements)
    set_aptos_font(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_aptos_font(cell.paragraphs)
    sanitized_name = promoter_name.replace(" ", "_")
    filename = f"{sanitized_name}_Declaration For Extension.docx"
    output_path = os.path.join(output_folder, filename)
    doc.save(output_path)
    return output_path

def update_FormB_File(promoter_name, project_name, office_address, extension_date, project_address):
    template_path = os.path.join('resources', 'FORM-B.docx')
    output_folder = os.path.join('generated', 'FORM_B')
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(template_path)
    extension_date_fmt = datetime.strptime(extension_date, "%Y-%m-%d").strftime("%d-%m-%Y")
    replacements = {'{{promoter_name}}': promoter_name, '{{project_name}}': project_name, '{{office_address}}': office_address, '{{extension_date}}': extension_date_fmt, '{{project_address}}': project_address}
    def replace_text_with_bold(paragraphs, replacements):
        for para in paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            if not any(ph in full_text for ph in replacements.keys()):
                continue
            new_runs = [(full_text, False)]
            for placeholder, value in replacements.items():
                updated_runs = []
                for text, bold in new_runs:
                    if placeholder in text:
                        parts = text.split(placeholder)
                        for i, part in enumerate(parts):
                            updated_runs.append((part, bold))
                            if i < len(parts) - 1:
                                updated_runs.append((value, True))
                    else:
                        updated_runs.append((text, bold))
                new_runs = updated_runs
            para.clear()
            for text, bold in new_runs:
                run = para.add_run(text)
                run.bold = bold
    replace_text_with_bold(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_with_bold(cell.paragraphs, replacements)
    set_aptos_font(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_aptos_font(cell.paragraphs)
    sanitized_name = promoter_name.replace(" ", "_")
    filename = f"{sanitized_name}_Form_B.docx"
    output_path = os.path.join(output_folder, filename)
    doc.save(output_path)
    return output_path

def update_FormatA_File(promoter_name, project_name, project_address, account_name, account_number, bank_name, branch_name, ifsc_code, date):
    template_path = os.path.join('resources', 'Format A.docx')
    output_folder = os.path.join('generated', 'Format_A')
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(template_path)
    date_fmt = datetime.strptime(date, "%Y-%m-%d").strftime("%d-%m-%Y")
    replacements = {'{{promoter_name}}': promoter_name, '{{project_name}}': project_name, '{{project_address}}': project_address, '{{account_name}}': account_name, '{{account_number}}': account_number, '{{bank_name}}': bank_name, '{{branch_name}}': branch_name, '{{ifsc_code}}': ifsc_code, '{{date}}': date_fmt}
    def replace_text_with_bold(paragraphs, replacements):
        for para in paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            if not any(ph in full_text for ph in replacements.keys()):
                continue
            new_runs = [(full_text, False)]
            for placeholder, value in replacements.items():
                updated_runs = []
                for text, bold in new_runs:
                    if placeholder in text:
                        parts = text.split(placeholder)
                        for i, part in enumerate(parts):
                            updated_runs.append((part, bold))
                            if i < len(parts) - 1:
                                updated_runs.append((value, True))
                    else:
                        updated_runs.append((text, bold))
                new_runs = updated_runs
            para.clear()
            for text, bold in new_runs:
                run = para.add_run(text)
                run.bold = bold
    replace_text_with_bold(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_with_bold(cell.paragraphs, replacements)
    set_aptos_font(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_aptos_font(cell.paragraphs)
    sanitized_name = promoter_name.replace(" ", "_")
    filename = f"{sanitized_name}_Format_A.docx"
    output_path = os.path.join(output_folder, filename)
    doc.save(output_path)
    return output_path

def update_FormatD_File(promoter_name, project_name, project_address, planning_authority, date):
    template_path = os.path.join('resources', 'FORMAT D.docx')
    output_folder = os.path.join('generated', 'Format_D')
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(template_path)
    date_fmt = datetime.strptime(date, "%Y-%m-%d").strftime("%d-%m-%Y")
    replacements = {'{{promoter_name}}': promoter_name, '{{project_name}}': project_name, '{{project_address}}': project_address, '{{planning_authority}}': planning_authority, '{{date}}': date_fmt}
    def replace_text_with_bold(paragraphs, replacements):
        for para in paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            if not any(ph in full_text for ph in replacements.keys()):
                continue
            new_runs = [(full_text, False)]
            for placeholder, value in replacements.items():
                updated_runs = []
                for text, bold in new_runs:
                    if placeholder in text:
                        parts = text.split(placeholder)
                        for i, part in enumerate(parts):
                            updated_runs.append((part, bold))
                            if i < len(parts) - 1:
                                updated_runs.append((value, True))
                    else:
                        updated_runs.append((text, bold))
                new_runs = updated_runs
            para.clear()
            for text, bold in new_runs:
                run = para.add_run(text)
                run.bold = bold
    replace_text_with_bold(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_with_bold(cell.paragraphs, replacements)
    set_aptos_font(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_aptos_font(cell.paragraphs)
    sanitized_name = promoter_name.replace(" ", "_")
    filename = f"{sanitized_name}_Format_D.docx"
    output_path = os.path.join(output_folder, filename)
    doc.save(output_path)
    return output_path

def update_Consent_Letter(promoter_name, office_address, project_name, registration_no, project_address, extension_date, date):
    template_path = os.path.join('resources', 'Consent Letter.docx')
    output_folder = os.path.join('generated', 'ConsentLetter')
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(template_path)
    date_fmt = datetime.strptime(date, "%Y-%m-%d").strftime("%d-%m-%Y")
    extension_date_fmt = datetime.strptime(extension_date, "%Y-%m-%d").strftime("%d-%m-%Y")
    replacements = {'{{promoter_name}}': promoter_name, '{{office_address}}': office_address, '{{project_name}}': project_name, '{{registration_no}}': registration_no, '{{project_address}}': project_address, '{{extension_date}}': extension_date_fmt, '{{date}}': date_fmt}
    def replace_text_with_bold(paragraphs, replacements):
        for para in paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            if not any(ph in full_text for ph in replacements.keys()):
                continue
            new_runs = [(full_text, False)]
            for placeholder, value in replacements.items():
                updated_runs = []
                for text, bold in new_runs:
                    if placeholder in text:
                        parts = text.split(placeholder)
                        for i, part in enumerate(parts):
                            updated_runs.append((part, bold))
                            if i < len(parts) - 1:
                                updated_runs.append((value, True))
                    else:
                        updated_runs.append((text, bold))
                new_runs = updated_runs
            para.clear()
            for text, bold in new_runs:
                run = para.add_run(text)
                run.bold = bold
    replace_text_with_bold(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_with_bold(cell.paragraphs, replacements)
    set_aptos_font(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_aptos_font(cell.paragraphs)
    sanitized_name = promoter_name.replace(" ", "_")
    filename = f"{sanitized_name}Consent Letter.docx"
    output_path = os.path.join(output_folder, filename)
    doc.save(output_path)
    return output_path

def update_form1(promoter_name, office_address, project_name, registration_no, as_on_date, date):
    template_path = os.path.join('resources', 'Form 1.docx')
    output_folder = os.path.join('generated', 'Form1')
    os.makedirs(output_folder, exist_ok=True)
    doc = Document(template_path)
    date_fmt = datetime.strptime(date, "%Y-%m-%d").strftime("%d-%m-%Y")
    replacements = {'{{promoter_name}}': promoter_name, '{{office_address}}': office_address, '{{project_name}}': project_name, '{{registration_no}}': registration_no, '{{as_on_date}}': as_on_date, '{{date}}': date_fmt}
    def replace_text_with_bold(paragraphs, replacements):
        for para in paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            if not any(ph in full_text for ph in replacements.keys()):
                continue
            new_runs = [(full_text, False)]
            for placeholder, value in replacements.items():
                updated_runs = []
                for text, bold in new_runs:
                    if placeholder in text:
                        parts = text.split(placeholder)
                        for i, part in enumerate(parts):
                            updated_runs.append((part, bold))
                            if i < len(parts) - 1:
                                updated_runs.append((value, True))
                    else:
                        updated_runs.append((text, bold))
                new_runs = updated_runs
            para.clear()
            for text, bold in new_runs:
                run = para.add_run(text)
                run.bold = bold
    replace_text_with_bold(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_with_bold(cell.paragraphs, replacements)
    set_aptos_font(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_aptos_font(cell.paragraphs)
    sanitized_name = promoter_name.replace(" ", "_")
    filename = f"{sanitized_name} Form1.docx"
    output_path = os.path.join(output_folder, filename)
    doc.save(output_path)
    return output_path


def update_form2a(promoter_name, office_address, project_name, registration_no, base_date, registration_date):
    import os
    from datetime import datetime
    from docx import Document

    template_path = os.path.join('resources', 'Form 2A.docx')
    output_folder = os.path.join('generated', 'form2a')
    os.makedirs(output_folder, exist_ok=True)

    # Parse base_date to determine financial year start for document naming
    base_date_dt = datetime.strptime(base_date, "%Y-%m-%d")
    if base_date_dt.month < 4:
        fy_start_year = base_date_dt.year - 1
    else:
        fy_start_year = base_date_dt.year
    financial_year_str = f"{fy_start_year}-{fy_start_year + 1}"

    # Parse registration_date (dd-mm-yyyy) to calculate cert_number
    reg_date_dt = datetime.strptime(registration_date, "%Y-%m-%d")

    today = datetime.today()
    current_fy_start_year = today.year if today.month >= 4 else today.year - 1

    # Calculate which financial year registration_date falls into relative to current FY
    cert_fy_start_year = current_fy_start_year
    fy_diff = 0
    while not (datetime(cert_fy_start_year, 4, 1) <= reg_date_dt <= datetime(cert_fy_start_year + 1, 3, 31)):
        cert_fy_start_year -= 1
        fy_diff += 1

    cert_number = f"{fy_diff + 1:02d}"  # "01" for current FY, then "02", "03"...

    # Load template
    doc = Document(template_path)

    replacements = {
        '{{promoter_name}}': promoter_name,
        '{{office_address}}': office_address,
        '{{project_name}}': project_name,
        '{{registration_no}}': registration_no,
        '{{date}}': today.strftime("%d-%m-%Y"),
        '{{cert_no_upd}}': cert_number,
        '{{financial_year}}': financial_year_str,
        '{{financial_year_date}}': financial_year_str
    }

    def replace_text_with_bold(paragraphs, replacements):
        for para in paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            if not any(ph in full_text for ph in replacements.keys()):
                continue
            new_runs = [(full_text, False)]
            for placeholder, value in replacements.items():
                updated_runs = []
                for text, bold in new_runs:
                    if placeholder in text:
                        parts = text.split(placeholder)
                        for i, part in enumerate(parts):
                            updated_runs.append((part, bold))
                            if i < len(parts) - 1:
                                updated_runs.append((value, True))
                    else:
                        updated_runs.append((text, bold))
                new_runs = updated_runs
            para.clear()
            for text, bold in new_runs:
                run = para.add_run(text)
                run.bold = bold

    replace_text_with_bold(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_with_bold(cell.paragraphs, replacements)

    set_aptos_font(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_aptos_font(cell.paragraphs)

    sanitized_name = promoter_name.replace(" ", "_")
    filename = f"{sanitized_name}_Form2A_{financial_year_str}.docx"
    output_path = os.path.join(output_folder, filename)
    doc.save(output_path)

    return [output_path]

# ============= DOCUMENT MAPPING FOR DYNAMIC UI ==============
DOCUMENTS = [
    {"key": "complaint_declaration", "name": "Complaint Declaration", "fn": "update_word_file", "fields": ["promoter_name", "project_name", "registration_no"]},
    {"key": "no_complaint_declaration", "name": "No Complaints Declaration", "fn": "generate_no_complaint_file", "fields": ["promoter_name", "project_name", "registration_no"]},
    {"key": "extension_application", "name": "Extension Application under Section 7(3)", "fn": "update_extension_application", "fields": ["registration_no", "extension_date", "promoter_name"]},
    {"key": "project_pert", "name": "Project PERT", "fn": "update_project_pert_file", "fields": ["project_name", "registration_no", "extension_date", "promoter_name"]},
    {"key": "cersai_declaration", "name": "CERSAI Declaration", "fn": "update_cersai_file", "fields": ["promoter_name", "project_name", "office_address", "project_address"]},
    {"key": "authorisation_letter", "name": "Authorization Letter", "fn": "update_authorization_letter", "fields": ["promoter_name", "project_name"]},
    {"key": "annexure_a", "name": "Annexure A", "fn": "update_annexure_a", "fields": ["promoter_name"]},
    {"key": "affidavit_reason_for_extension", "name": "Affidavit Reason for Extension", "fn": "update_AffidavitReason_for_Extension", "fields": ["promoter_name", "project_name", "registration_no"]},
    {"key": "consent_extension_tabular", "name": "Consent for Extension (Tabular)", "fn": "update_consent_Extension_Tabular", "fields": ["promoter_name", "project_name", "registration_no", "extension_date"]},
    {"key": "declaration_extension", "name": "Declaration For Extension", "fn": "update_Declaration_For_Extension", "fields": ["promoter_name", "project_name", "registration_no", "extension_date"]},
    {"key": "form_b", "name": "Form B", "fn": "update_FormB_File", "fields": ["promoter_name", "project_name", "office_address", "extension_date", "project_address"]},
    {"key": "format_a", "name": "Format A", "fn": "update_FormatA_File", "fields": ["promoter_name", "project_name", "project_address", "account_name", "account_number", "bank_name", "branch_name", "ifsc_code"]},
    {"key": "format_d", "name": "Format D", "fn": "update_FormatD_File", "fields": ["promoter_name", "project_name", "project_address", "planning_authority"]},
    {"key": "consent_letter", "name": "Consent Letter", "fn": "update_Consent_Letter", "fields": ["promoter_name", "office_address", "project_name", "registration_no", "project_address", "extension_date"]},
    {"key": "form1", "name": "Form 1", "fn": "update_form1", "fields": ["promoter_name", "office_address", "project_name", "registration_no", "as_on_date"]},
    {"key": "form2a", "name": "Form 2A (for all required years)", "fn": "update_form2a", "fields": ["promoter_name", "office_address", "project_name", "registration_no", "registration_date"]}
]
DOC_FN_MAP = {d["key"]: d["fn"] for d in DOCUMENTS}
DOC_FIELD_MAP = {d["key"]: d["fields"] for d in DOCUMENTS}

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        docs_selected = request.form.getlist('documents')
        all_fields = set()
        for doc in docs_selected:
            all_fields |= set(DOC_FIELD_MAP[doc])
        data = {field: request.form.get(field, "") for field in all_fields}
        generated_files = []
        for doc in docs_selected:
            fn = globals()[DOC_FN_MAP[doc]]
            fields = DOC_FIELD_MAP[doc]
            if doc == 'form2a':
                now = datetime.now()
                fy_start_year = now.year if now.month >= 4 else now.year - 1
                base_date = f"{fy_start_year}-04-01"  # e.g., for financial year start
                registration_date = data.get('registration_date', '')  # get registration_date from user input
                args = [data.get(f, "") for f in fields if f != 'registration_date'] + [base_date, registration_date]
            elif doc == 'cersai_declaration':
                today_str = datetime.now().strftime("%Y-%m-%d")
                args = [data.get(f, "") for f in fields] + [today_str]
            elif doc == 'no_complaint_declaration':
                today_str = datetime.now().strftime("%Y-%m-%d")
                args = [data.get(f, "") for f in fields] + [today_str]
            elif doc == 'extension_application':
                today_str = datetime.now().strftime("%Y-%m-%d")
                args = [data.get(f, "") for f in fields] + [today_str]
            elif doc == 'authorisation_letter':
                today_str = datetime.now().strftime("%Y-%m-%d")
                args = [data.get(f, "") for f in fields] + [today_str]
            elif doc == 'annexure_a':
                today_str = datetime.now().strftime("%Y-%m-%d")
                args = [data.get(f, "") for f in fields] + [today_str]
            elif doc == 'format_d':
                today_str = datetime.now().strftime("%Y-%m-%d")
                args = [data.get(f, "") for f in fields] + [today_str]
            elif doc == 'form1':
                today_str = datetime.now().strftime("%Y-%m-%d")
                args = [data.get(f, "") for f in fields] + [today_str]
            else:
                args = [data.get(f, "") for f in fields]
            result = fn(*args)
            if doc == 'form2a' and isinstance(result, list):
                generated_files.extend(result)
            else:
                generated_files.append(result)
        temp_zip = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
        with zipfile.ZipFile(temp_zip.name, 'w') as zipf:
            for file_path in generated_files:
                zipf.write(file_path, os.path.basename(file_path))
        zip_filename = f"Documents_{datetime.now().strftime('%Y%m%d%H%M%S')}.zip"
        return send_file(temp_zip.name, as_attachment=True, download_name=zip_filename)
    else:
        doc_inputs_json = json.dumps(DOC_FIELD_MAP)
        sorted_documents = sorted(DOCUMENTS, key=lambda d: d['name'].lower())
        return render_template('index.html', documents=sorted_documents, doc_inputs_json=doc_inputs_json)

if __name__ == '__main__':
    app.run(debug=True)
